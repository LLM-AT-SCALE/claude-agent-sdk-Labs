"""Reading a loan application out of the file the user uploaded.

A loan application arrives as a PDF or a Word document, not as plain text, so
the lab has to open both. Dispatch is by extension, and every parser returns
the same thing: the narrative as a single string.

No OCR. A scanned or photographed PDF has no text layer, and rather than
guessing at pixels the lab says so plainly and stops.
"""

from __future__ import annotations

import io
from pathlib import PurePosixPath

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt")

# Below this, whatever we pulled out is not a loan application narrative —
# almost always a scanned PDF with no text layer behind the image.
MIN_USABLE_CHARS = 200


class UnreadableDocument(Exception):
    """The file opened but produced nothing worth scoring."""


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - surfaced to the student
        raise UnreadableDocument(
            "That PDF could not be opened. It may be corrupt or password-protected."
        ) from exc

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:  # noqa: BLE001
            raise UnreadableDocument(
                "That PDF is password-protected. Please upload an unlocked copy."
            ) from exc

    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - one bad page must not sink the file
            pages.append("")

    return "\n\n".join(pages)


def _from_docx(data: bytes) -> str:
    from docx import Document

    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise UnreadableDocument(
            "That Word file could not be opened. It may be corrupt, or saved in the "
            "older .doc format, which this lab does not read."
        ) from exc

    blocks = [p.text for p in document.paragraphs if p.text and p.text.strip()]

    # Loan narratives routinely put the financials in a table. Skipping tables
    # would silently lose the very figures DSCR and LTV depend on.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            deduped: list[str] = []
            for cell in cells:
                if not deduped or deduped[-1] != cell:  # merged spans repeat
                    deduped.append(cell)
            line = " | ".join(c for c in deduped if c)
            if line:
                blocks.append(line)

    return "\n\n".join(blocks)


def _from_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


_PARSERS = {".pdf": _from_pdf, ".docx": _from_docx, ".txt": _from_txt}


def extract_text(filename: str, data: bytes) -> str:
    """Return the narrative held in ``data``.

    Raises UnreadableDocument for an unsupported format, an empty file, or a
    PDF with no text layer.
    """
    if not data:
        raise UnreadableDocument("That file is empty.")

    suffix = PurePosixPath(filename.replace("\\", "/")).suffix.lower()
    parser = _PARSERS.get(suffix)

    if parser is None:
        raise UnreadableDocument(
            f"'{filename}' is not a supported format. "
            f"Upload a {', '.join(SUPPORTED_EXTENSIONS)} file."
        )

    text = parser(data).strip()

    if len(text) < MIN_USABLE_CHARS:
        raise UnreadableDocument(
            "No readable text was found in that document. If it is a scanned or "
            "photographed PDF, this lab cannot read it — there is no OCR step. "
            "Please upload a text-based PDF, a .docx, or a .txt file."
        )

    return text
